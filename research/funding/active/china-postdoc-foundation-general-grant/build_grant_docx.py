import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import latex2mathml.converter
import mathml2omml

# -------------------------------------------------------------
# XML / Styling Helpers
# -------------------------------------------------------------

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    kwargs: top, bottom, left, right
    value: dict(sz=12, val='single', color='000000')
    """
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, opts in kwargs.items():
        edge_data = OxmlElement(f'w:{edge}')
        for key, val in opts.items():
            edge_data.set(qn(f'w:{key}'), str(val))
        tcBorders.append(edge_data)
    tcPr.append(tcBorders)

def set_cell_background(cell, fill_hex="F2F2F2"):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_run_font(run, font_name="宋体", ascii_font="Times New Roman", size_pt=12, bold=False, italic=False, color=None):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), ascii_font)
    rPr.append(rFonts)

def latex_to_omml(latex_str, is_display=False):
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        omml = mathml2omml.convert(mathml)
        if 'xmlns:m=' not in omml:
            omml = omml.replace('<m:oMath', '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"', 1)
        if is_display:
            return f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</m:oMathPara>'
        return omml
    except Exception as e:
        print(f"Warning: LaTeX conversion failed for '{latex_str}': {e}")
        return None

def clean_markdown_links(text):
    # [[target|display]] -> display
    text = re.sub(r'\[\[(?:[^|\]]*\|)?([^\]]+)\]\]', r'\1', text)
    # [display](url) -> display
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

def parse_and_add_text_runs(p, text, default_font="宋体", default_ascii="Times New Roman", default_size=12, default_bold=False, default_italic=False, default_color=None):
    """
    Parses inline elements: LaTeX $...$, bold **...**, italic *...*, code `...`, and regular text
    """
    # Regex to tokenize: $...$, **...**, *...*, `...`
    pattern = re.compile(r'(\$[^\$]+\$|\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)')
    tokens = pattern.split(text)

    for token in tokens:
        if not token:
            continue
        if token.startswith('$') and token.endswith('$') and len(token) > 2:
            latex_code = token[1:-1]
            omml = latex_to_omml(latex_code, is_display=False)
            if omml:
                try:
                    p._element.append(parse_xml(omml))
                    continue
                except Exception as e:
                    pass
            # Fallback to plain text
            run = p.add_run(latex_code)
            set_run_font(run, default_font, default_ascii, default_size, default_bold, italic=True, color=default_color)
        elif token.startswith('**') and token.endswith('**') and len(token) >= 4:
            inner = token[2:-2]
            # Recursively handle inner if it contains math or code
            sub_tokens = re.compile(r'(\$[^\$]+\$|`[^`]+`)').split(inner)
            for st in sub_tokens:
                if not st:
                    continue
                if st.startswith('$') and st.endswith('$') and len(st) > 2:
                    latex_code = st[1:-1]
                    omml = latex_to_omml(latex_code, is_display=False)
                    if omml:
                        try:
                            p._element.append(parse_xml(omml))
                            continue
                        except Exception:
                            pass
                    run = p.add_run(latex_code)
                    set_run_font(run, default_font, default_ascii, default_size, bold=True, italic=True, color=default_color)
                elif st.startswith('`') and st.endswith('`') and len(st) >= 2:
                    code_text = st[1:-1]
                    run = p.add_run(code_text)
                    set_run_font(run, "Consolas", "Consolas", default_size - 0.5, bold=True, color=default_color)
                else:
                    run = p.add_run(st)
                    set_run_font(run, "黑体" if default_font == "宋体" else default_font, default_ascii, default_size, bold=True, italic=default_italic, color=default_color)
        elif token.startswith('*') and token.endswith('*') and len(token) >= 3 and not token.startswith('**'):
            inner = token[1:-1]
            run = p.add_run(inner)
            set_run_font(run, default_font, default_ascii, default_size, default_bold, italic=True, color=default_color)
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            code_text = token[1:-1]
            run = p.add_run(code_text)
            set_run_font(run, "Consolas", "Consolas", default_size - 0.5, default_bold, default_italic, color=default_color)
        else:
            run = p.add_run(token)
            set_run_font(run, default_font, default_ascii, default_size, default_bold, default_italic, color=default_color)

def init_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    return doc

# -------------------------------------------------------------
# Document Generators
# -------------------------------------------------------------

def build_full_draft_docx(md_path, output_path, assets_dir):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = init_doc()

    # 1. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(18)
    p_title.paragraph_format.line_spacing = 1.3
    run_title = p_title.add_run("中国博士后科学基金第 80 批面上资助申请书正文")
    set_run_font(run_title, font_name="黑体", ascii_font="Times New Roman", size_pt=18, bold=True)

    lines = content.split('\n')
    i = 0
    in_yaml = False
    in_section_3_meta = False

    while i < len(lines):
        line = lines[i].strip()

        # Skip YAML
        if i == 0 and line == '---':
            in_yaml = True
            i += 1
            continue
        if in_yaml:
            if line == '---':
                in_yaml = False
            i += 1
            continue

        # Skip section 3 (internal verification meta)
        if line.startswith('## 三、正文版本与核验边界'):
            in_section_3_meta = True
            i += 1
            continue
        if in_section_3_meta:
            i += 1
            continue

        # Skip main markdown title as we already added styled title
        if line.startswith('# 中国博士后科学基金'):
            i += 1
            continue

        # Skip status indicator lines
        if line.startswith('**状态：') or line.startswith('**状态:'):
            i += 1
            continue

        # Empty line
        if not line:
            i += 1
            continue

        # Skip internal verification notes
        if line.startswith('> **草稿核验资料') or '（不导出至申请书）' in line:
            # Skip until blockquote ends
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                i += 1
            continue

        # Blockquote (Callout)
        if line.startswith('>'):
            callout_text = line[1:].strip()
            # Clean links
            callout_text = clean_markdown_links(callout_text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2

            # Style as note
            parse_and_add_text_runs(p, callout_text, default_font="楷体", default_size=10.5, default_color=RGBColor(80, 80, 80))
            i += 1
            continue

        # Level 1 Heading (## 一、项目基本信息, ## 二、项目信息)
        if line.startswith('## '):
            h_text = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            set_run_font(run, font_name="黑体", ascii_font="Times New Roman", size_pt=14, bold=True)
            i += 1
            continue

        # Level 3 Heading (#### （一）国内外研究现状与主要不足...)
        if line.startswith('#### '):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p, f"**{h_text}**", default_font="黑体", default_size=12, default_bold=True)
            i += 1
            continue

        # Reference heading (**参考文献...**)
        if line.startswith('**参考文献'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p, "**参考文献：**", default_font="黑体", default_size=10.5, default_bold=True)
            i += 1
            continue

        # Numbered reference entry ([1] Kronbichler...)
        if re.match(r'^\[\d+\]', line):
            cleaned_line = clean_markdown_links(line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)
            parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=10)
            i += 1
            continue

        # Level 2 Heading (### 1. 选题依据...)
        if line.startswith('### '):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            set_run_font(run, font_name="黑体", ascii_font="Times New Roman", size_pt=12.5, bold=True)
            i += 1
            continue

        # Table detection
        if line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            # Collect table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            raw_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if not any('---' in c for c in cells):
                    raw_rows.append(cells)

            if raw_rows:
                tbl = doc.add_table(rows=len(raw_rows), cols=len(raw_rows[0]))
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False

                # Column widths: 3.5cm, 8.5cm, 3.5cm
                col_widths = [Cm(3.5), Cm(8.6), Cm(3.5)]
                for r_idx, row_data in enumerate(raw_rows):
                    row = tbl.rows[r_idx]
                    is_header = (r_idx == 0)
                    for c_idx, cell_text in enumerate(row_data):
                        cell = row.cells[c_idx]
                        if c_idx < len(col_widths):
                            cell.width = col_widths[c_idx]
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                        # 3-line table borders
                        border_opts = {}
                        if is_header:
                            border_opts['top'] = dict(sz=12, val='single', color='000000')
                            border_opts['bottom'] = dict(sz=6, val='single', color='000000')
                            set_cell_background(cell, "F2F4F7")
                        else:
                            if r_idx == len(raw_rows) - 1:
                                border_opts['bottom'] = dict(sz=12, val='single', color='000000')
                            else:
                                border_opts['bottom'] = dict(sz=2, val='single', color='E0E0E0')

                        set_cell_border(cell, **border_opts)

                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_header or c_idx == 0 or c_idx == 2) else WD_ALIGN_PARAGRAPH.LEFT
                        cp.paragraph_format.space_before = Pt(2)
                        cp.paragraph_format.space_after = Pt(2)
                        cp.paragraph_format.line_spacing = 1.15

                        cleaned_cell_text = clean_markdown_links(cell_text)
                        parse_and_add_text_runs(cp, cleaned_cell_text, default_font="宋体", default_size=10.5, default_bold=is_header)

                # Space after table
                p_spacer = doc.add_paragraph()
                p_spacer.paragraph_format.space_before = Pt(0)
                p_spacer.paragraph_format.space_after = Pt(6)
            continue

        # Image detection ![alt](src)
        img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_src = img_match.group(2)

            # Resolve image path
            img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), img_src))
            if not os.path.exists(img_path) and 'assets/' in img_src:
                img_path = os.path.join(assets_dir, os.path.basename(img_src))

            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(4)
                # Width: 15.2 cm
                p_img.add_run().add_picture(img_path, width=Cm(15.2))
            i += 1
            continue

        # Image caption <center><b>...</b></center>
        if '<center>' in line and '<b>' in line:
            caption = re.sub(r'<[^>]+>', '', line).strip()
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(8)
            p_cap.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p_cap, caption, default_font="黑体", default_size=10.5, default_bold=True)
            i += 1
            continue

        # Display math $$ ... $$
        if line.startswith('$$'):
            # Collect multiline display math if needed
            math_lines = []
            if line == '$$':
                i += 1
                while i < len(lines) and lines[i].strip() != '$$':
                    math_lines.append(lines[i].strip())
                    i += 1
                i += 1 # skip closing $$
            else:
                math_lines.append(line.replace('$$', '').strip())
                i += 1

            math_code = ' '.join(math_lines)
            omml_para = latex_to_omml(math_code, is_display=True)
            p_math = doc.add_paragraph()
            p_math.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_math.paragraph_format.space_before = Pt(6)
            p_math.paragraph_format.space_after = Pt(6)
            if omml_para:
                try:
                    p_math._element.append(parse_xml(omml_para))
                except Exception as e:
                    print(f"Error appending display math: {e}")
                    run = p_math.add_run(math_code)
                    set_run_font(run, "Times New Roman", "Times New Roman", 11, italic=True)
            continue

        # Sub-heading (e.g. **一、 申请人前期研究积累与核心技术储备**)
        if line.startswith('**一、') or line.startswith('**二、') or line.startswith('**三、') or line.startswith('**总体思路：**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Pt(24)

            cleaned_line = clean_markdown_links(line)
            parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=12)
            i += 1
            continue

        # Regular Body Paragraph or Numbered Item
        cleaned_line = clean_markdown_links(line)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Pt(24) # 2 Chinese characters indent

        parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=12)
        i += 1

    doc.save(output_path)
    print(f"Generated full draft: {output_path}")

def build_part1_5_anonymous_docx(md_path, output_path, assets_dir):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = init_doc()

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(18)
    p_title.paragraph_format.line_spacing = 1.3
    run_title = p_title.add_run("中国博士后科学基金第 80 批面上资助项目信息（1–5 部分）")
    set_run_font(run_title, font_name="黑体", ascii_font="Times New Roman", size_pt=16, bold=True)

    lines = content.split('\n')
    i = 0
    in_target_section = False

    while i < len(lines):
        line = lines[i].strip()

        # We start capturing from ### 1. 选题依据
        if line.startswith('### 1. 选题依据'):
            in_target_section = True

        # Stop capturing before ### 6. 研究基础
        if line.startswith('### 6. 研究基础'):
            in_target_section = False
            break

        if not in_target_section:
            i += 1
            continue

        # Skip status lines
        if line.startswith('**状态：') or line.startswith('**状态:'):
            i += 1
            continue

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip internal verification notes
        if line.startswith('> **草稿核验资料') or '（不导出至申请书）' in line:
            # Skip until blockquote ends
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                i += 1
            continue

        # Level 3 Heading (#### （一）国内外研究现状与主要不足...)
        if line.startswith('#### '):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p, f"**{h_text}**", default_font="黑体", default_size=12, default_bold=True)
            i += 1
            continue

        # Reference heading (**参考文献...**)
        if line.startswith('**参考文献'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p, "**参考文献：**", default_font="黑体", default_size=10.5, default_bold=True)
            i += 1
            continue

        # Numbered reference entry ([1] Kronbichler...)
        if re.match(r'^\[\d+\]', line):
            cleaned_line = clean_markdown_links(line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)
            parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=10)
            i += 1
            continue

        # Level 2 Heading (### 1. 选题依据...)
        if line.startswith('### '):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            set_run_font(run, font_name="黑体", ascii_font="Times New Roman", size_pt=13, bold=True)
            i += 1
            continue

        # Image detection ![alt](src)
        img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img_match:
            img_src = img_match.group(2)
            img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), img_src))
            if not os.path.exists(img_path) and 'assets/' in img_src:
                img_path = os.path.join(assets_dir, os.path.basename(img_src))

            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(4)
                p_img.add_run().add_picture(img_path, width=Cm(15.2))
            i += 1
            continue

        # Image caption <center><b>...</b></center>
        if '<center>' in line and '<b>' in line:
            caption = re.sub(r'<[^>]+>', '', line).strip()
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(8)
            p_cap.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p_cap, caption, default_font="黑体", default_size=10.5, default_bold=True)
            i += 1
            continue

        # Display math $$ ... $$
        if line.startswith('$$'):
            math_lines = []
            if line == '$$':
                i += 1
                while i < len(lines) and lines[i].strip() != '$$':
                    math_lines.append(lines[i].strip())
                    i += 1
                i += 1
            else:
                math_lines.append(line.replace('$$', '').strip())
                i += 1

            math_code = ' '.join(math_lines)
            omml_para = latex_to_omml(math_code, is_display=True)
            p_math = doc.add_paragraph()
            p_math.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_math.paragraph_format.space_before = Pt(6)
            p_math.paragraph_format.space_after = Pt(6)
            if omml_para:
                try:
                    p_math._element.append(parse_xml(omml_para))
                except Exception as e:
                    print(f"Error appending display math: {e}")
                    run = p_math.add_run(math_code)
                    set_run_font(run, "Times New Roman", "Times New Roman", 11, italic=True)
            continue

        # Regular Body Paragraph
        cleaned_line = clean_markdown_links(line)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Pt(24) # 2 Chinese chars indent

        parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=12)
        i += 1

    doc.save(output_path)
    print(f"Generated anonymous part 1-5: {output_path}")

def build_part6_research_basis_docx(md_path, output_path, assets_dir):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = init_doc()

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(18)
    p_title.paragraph_format.line_spacing = 1.3
    run_title = p_title.add_run("中国博士后科学基金第 80 批面上资助项目信息（6、研究基础）")
    set_run_font(run_title, font_name="黑体", ascii_font="Times New Roman", size_pt=16, bold=True)

    lines = content.split('\n')
    i = 0
    in_target_section = False

    while i < len(lines):
        line = lines[i].strip()

        # Start capturing from ### 6. 研究基础
        if line.startswith('### 6. 研究基础'):
            in_target_section = True
            i += 1
            continue

        # Stop capturing before ## 三、正文版本与核验边界
        if line.startswith('## 三、正文版本与核验边界'):
            in_target_section = False
            break

        if not in_target_section:
            i += 1
            continue

        # Skip status lines
        if line.startswith('**状态：') or line.startswith('**状态:'):
            i += 1
            continue

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Blockquote (Figure explanation, etc.)
        if line.startswith('>'):
            callout_text = line[1:].strip()
            callout_text = clean_markdown_links(callout_text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2

            parse_and_add_text_runs(p, callout_text, default_font="楷体", default_size=10.5, default_color=RGBColor(60, 60, 60))
            i += 1
            continue

        # Sub-heading (e.g. **一、 申请人前期研究积累与核心技术储备**)
        if line.startswith('**一、') or line.startswith('**二、') or line.startswith('**三、'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.keep_with_next = True

            cleaned_line = clean_markdown_links(line)
            parse_and_add_text_runs(p, cleaned_line, default_font="黑体", default_size=12.5, default_bold=True)
            i += 1
            continue

        # Image detection ![alt](src)
        img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img_match:
            img_src = img_match.group(2)
            img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), img_src))
            if not os.path.exists(img_path) and 'assets/' in img_src:
                img_path = os.path.join(assets_dir, os.path.basename(img_src))

            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(4)
                p_img.add_run().add_picture(img_path, width=Cm(15.2))
            i += 1
            continue

        # Image caption <center><b>...</b></center>
        if '<center>' in line and '<b>' in line:
            caption = re.sub(r'<[^>]+>', '', line).strip()
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(6)
            p_cap.paragraph_format.keep_with_next = True
            parse_and_add_text_runs(p_cap, caption, default_font="黑体", default_size=10.5, default_bold=True)
            i += 1
            continue

        # Regular Body Paragraph
        cleaned_line = clean_markdown_links(line)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Pt(24) # 2 Chinese chars indent

        parse_and_add_text_runs(p, cleaned_line, default_font="宋体", default_size=12)
        i += 1

    doc.save(output_path)
    print(f"Generated research basis part 6: {output_path}")

if __name__ == '__main__':
    base_dir = r"C:\workspace\dut-postdoc\research\funding\active\china-postdoc-foundation-general-grant"
    md_file = os.path.join(base_dir, "80th-2026-application-draft.md")
    assets = os.path.join(base_dir, "assets")

    out_full = os.path.join(base_dir, "80th-2026-application-draft-完整审阅版.docx")
    out_p1_5 = os.path.join(base_dir, "80th-2026-项目信息(1-5部分)-审阅版.docx")
    out_p6 = os.path.join(base_dir, "80th-2026-项目信息(6部分)-研究基础.docx")

    build_full_draft_docx(md_file, out_full, assets)
    build_part1_5_anonymous_docx(md_file, out_p1_5, assets)
    build_part6_research_basis_docx(md_file, out_p6, assets)
    print("All 3 DOCX files generated successfully!")
